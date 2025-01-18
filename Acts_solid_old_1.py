from tkinter.ttk import Combobox
#import numpy as np
from tkinter import *
from tkinter import filedialog
import pandas as pd
from docxtpl import DocxTemplate
from tkinter.messagebox import showinfo
from tkinter import font
import sys
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


#при запуске exe'шника все внутренности программы распаковываюся
# во временную папку Windows. Соответственно, в скрипте нужно обращаться к ним
#Для обращения к этим файлам и нужна функция ниже
#https://ru.stackoverflow.com/questions/1472473/%d0%90%d0%b2%d1%82%d0%be%d0%bd%d0%be%d0%bc%d0%bd%d1%8b%d0%b9-%d0%b8%d1%81%d0%bf%d0%be%d0%bb%d0%bd%d1%8f%d0%b5%d0%bc%d1%8b%d0%b9-%d1%84%d0%b0%d0%b9%d0%bb-%d0%b2-python
# По ссылке смотри описание
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Функция вызывающее информационное окно для напоминания о выборе нужного файла спецификации
def open_spec():
    showinfo(title="Информация", message="Выберете спецификацию в формате xlsx")

# Функция вызывающее информационное окно для напоминания о выборе нужного файла спецификации
def info_head():
    showinfo(title="Информация", message="Выберете файл с полями акта в формате xlsx")

# Функция вызывающая информационное окно когда не загружена спецификация при попытке создания акта
def make_act():
    showinfo(title="Информация", message="Проверьте что спецификация загружена!")

# Функция создания акта
def safe_act():
    # Проверка если спецификация была не загружена, то нет возможности создать акт
    if len(tube_list) <= 1:
        make_act()
        return
    global document

    # Данные для заполнения шаблона
    context = {
        'station': station.get(),
        'calc': calc.get(),
        'company': company.get(),
        'object': obj.get(),
        'address': address.get(),
        # 'number': number.get(),
        # 'name': name.get(),
        'data': data.get()
    }

    if type_choies.get() == 'входного контроля элементов трубопровода':
        context['number'] = '2'
        context['name'] = 'ВХОДНОГО КОНТРОЛЯ ЭЛЕМЕНТОВ ТРУБОПРОВОДА'
    elif type_choies.get() == 'входного контроля материалов':
        context['number'] = '3'
        context['name'] = 'ВХОДНОГО КОНТРОЛЯ МАТЕРИАЛОВ'
    elif type_choies.get() == 'проверки чистоты труб':
        context['number'] = '4'
        context['name'] = 'ПРОВЕРКИ ЧИСТОТЫ ТРУБ'
    elif type_choies.get() == 'монтажа опор под трубопроводы':
        context['number'] = '5'
        context['name'] = 'МОНТАЖА ОПОР ПОД ТРУБОПРОВОДЫ'
    elif type_choies.get() == 'монтажа трубопроводов':
        context['number'] = '6'
        context['name'] = 'МОНТАЖА ТРУБОПРОВОДОВ'

    # Если спецификация загружена и выбран файл для сохранения акта, то выполняется код ниже
    filepath = filedialog.asksaveasfilename(defaultextension='docx', initialfile='Акт '+str(context['name']).lower())
    if filepath != "" and len(tube_list)>1:
        # Заголовки таблицы (при использовании шаблона они не нужны и используюся только для подсчета
        # столбцов
        headers = ('№ ', 'Наименование', 'Размеры, материал',
                   'Техническая характеристика', 'Кол-во', 'Ед.')

        # Заполнение шаблона данными
        document.render(context)

        # Получение списка таблиц из файла шаблона
        all_tables = document.tables
        # Поиск таблицы с одной строкой в шаблоне
        new_table = all_tables[0]
        # Количество колонок таблицы
        cols_number = len(headers)
        #Если выбран тип акта основное оборудование
        if type_choies.get() == 'входного контроля элементов трубопровода':

            new_elements_list=elements_list.copy()
            # Заполняется таблица шаблона списком основного оборудования
            for row in new_elements_list:
                row_cells = new_table.add_row().cells
                for i in range(cols_number):
                    row_cells[i].text = str(row[i])
                    # центруем ячейки в таблице
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Далее аналогично двум другим выборкам
        elif type_choies.get() =='входного контроля материалов':

            new_pad_list=pad_list.copy()
            # Заполняется таблица шаблона списком основного оборудования
            for row in new_pad_list:
                row_cells = new_table.add_row().cells
                for i in range(cols_number):
                    row_cells[i].text = str(row[i])
                    # центруем ячейки в таблице
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif type_choies.get() == 'проверки чистоты труб':
            new1_tube_list=tube_list.copy()
            # Заполняется таблица шаблона списком основного оборудования
            for row in new1_tube_list:
                row_cells = new_table.add_row().cells
                for i in range(cols_number):
                    row_cells[i].text = str(row[i])
                    # центруем ячейки в таблице
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



        elif type_choies.get() =='монтажа опор под трубопроводы':

            ne_support_list=support_list.copy()
            # Заполняется таблица шаблона списком основного оборудования
            for row in ne_support_list:
                row_cells = new_table.add_row().cells
                for i in range(cols_number):
                    row_cells[i].text = str(row[i])
                    # центруем ячейки в таблице
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif type_choies.get() == 'монтажа трубопроводов':


            new_tube_list=tube_list.copy()
            # Заполняется таблица шаблона списком основного оборудования
            for row in new_tube_list:
                row_cells = new_table.add_row().cells
                for i in range(cols_number):
                    row_cells[i].text = str(row[i])
                    # центруем ячейки в таблице
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



        # Если выборки не сделано, то таблица в шаблоне наполняется полным списком из спецификации
        else:
            print('ok')

    document.save(filepath)

# функция открытия файла с полями шапки
def open_head():
    # Информационное окно напоминающее, что нужно выбрать для открытия файл спецификации формата xlsx
    info_head()
    head_path=filedialog.askopenfilename()
    global name_station
    global number_calc
    global company_name
    global obj_name
    global addr_obj
    global dt
    # Если файл выбран, то данные из листа с названием Table 1 читаются датафрейм пандас
    if head_path !="":
       df = pd.read_excel(head_path)
       name_station.set(df[df.columns[1]].iloc[0])
       number_calc.set(df[df.columns[1]].iloc[1])
       company_name.set(df[df.columns[1]].iloc[2])
       obj_name.set(df[df.columns[1]].iloc[3])
       addr_obj.set(df[df.columns[1]].iloc[4])
       dt.set(df[df.columns[1]].iloc[5])


#Функция открытия xlsx файла с помощью диалогового окна выбора файла в системе
def open_table():
    # Информационное окно напоминающее, что нужно выбрать для открытия файл спецификации формата xlsx
    open_spec()
    table_path=filedialog.askopenfilename()

    # Если файл выбран, то данные из листа с названием Table 1 читаются датафрейм пандас
    if table_path !="":
        global model_dict
        #df = pd.read_excel(table_path, sheet_name='Table 1', skiprows=2)
        df = pd.read_excel(table_path)
        # удаляем лишние столбцы из исходного датафрейма
        df = df.drop(df.columns[[0, 1, 3, 5, 6, 7, 11, 12, 15, 16, 17, 18]], axis=1)
        # cоздаем список из названий столбцов полученного выше датафрейм
        columns_names_df = df.columns.tolist()
        # удаляем все строки которые имеют Nan в первом стоблце датафрейм
        df_clean = df.dropna(subset=columns_names_df[0])
        # создаем список номеров из словаря Котенко для труб
        list_of_tube = list(range(7, 13))

        # создаем датафрейм df_tube в котором только трубы
        df_tube = df_clean[df_clean[columns_names_df[1]].isin(list_of_tube)]

        # создадим пустой датафрейм
        df_tube_new = pd.DataFrame(columns=columns_names_df)
        # добави в него 1 строку с числами
        df_tube_new.loc[0]=range(0, len(columns_names_df))

        # список уникальных значений в 2 третьем столбце df_tube
        uniq_val_tube = df_tube[columns_names_df[1]].unique()
        # проходимся по циклу через все уникальные значения 2 столбца датафрейма труб
        for i in range(0, len(uniq_val_tube)):
            df_tube_i = df_tube[df_tube[columns_names_df[1]] == uniq_val_tube[i]]
            # список уникальных значений в 3 третьем столбце для каждого уникального во 2 столбце
            uniq_val_tube_i = df_tube_i[columns_names_df[2]].unique()
            # проходимся по циклу через все уникальные значения в третьем столбце каждого уникального
            #во втором
            for j in range(0,len(uniq_val_tube_i)):
                df_tube_j = df_tube_i[df_tube_i[columns_names_df[2]] == uniq_val_tube_i[j]]
                # cчитаем сумму длинн в 6 столбце полученного датафрейма df_tube_j и оставляем 2 знака после
                #запятой
                sum_j = round(df_tube_j[columns_names_df[5]].sum(axis=0), 2)
                # удаляем дубликаты то есть по сути все строки кроме 1 строки в df_tube_j
                df_tube_j = df_tube_j.drop_duplicates(subset=[columns_names_df[1]])
                # заносим в первую и единственную строку в 6 столбец значение sum_j
                df_tube_j.iat[0, 5] = sum_j
                # в 4 столбце удаляем значение длины участка в скобках которое было в xlsx файле
                df_tube_j.iat[0, 3] = df_tube_j.iat[0, 3][:df_tube_j.iat[0, 3].find('(')]
                # в 6 столбце меняем м на п.м
                df_tube_j.iat[0, 6]='п.м.'

                df_tube_new = pd.concat([df_tube_new, df_tube_j], axis=0)  # объединение по вертикали

        # удаляем лишние столбцы из  датафрейма df_tube_new
        df_tube_new = df_tube_new.drop(df.columns[[1, 4]], axis=1)

        # удаляем первую строку из датафрейма df_tube_new
        df_tube_new=df_tube_new.iloc[1:]


        #из датафрейма df создаем двумерный массив
        tube_arr = df_tube_new.to_numpy()

        global tube_list
        #создаем список
        tube_list=tube_arr.tolist()
        #print(tube_list)

        # создаем столбец сквозной нумерации
        j = 1
        for i in range(0, len(tube_list)):
            tube_list[i].insert(0, j)
            j += 1

        # номер соотвествующий прокладке
        numb_pad = 29
        # создаем датафрейм df_pad в котором только прокладки
        df_pad = df_clean[df_clean[columns_names_df[1]] == numb_pad]
        # удаляем лишние столбцы из  датафрейма df_pad
        df_pad = df_pad.drop(df.columns[[1, 5]], axis=1)
        #из датафрейма df создаем двумерный массив
        pad_arr = df_pad.to_numpy()

        global pad_list
        # создаем список
        pad_list = pad_arr.tolist()

        # создаем столбец сквозной нумерации
        j = 1
        for i in range(0, len(pad_list)):
            pad_list[i].insert(0, j)
            j += 1

        # создаем список номеров из словаря Котенко для опор
        list_of_support = list(range(68, 70))
        # создаем датафрейм df_support в котором только опоры и хомуты
        df_support = df_clean[df_clean[columns_names_df[1]].isin(list_of_support)]

        # удаляем лишние столбцы из  датафрейма df_support
        df_support = df_support.drop(df.columns[[1, 5]], axis=1)
        #из датафрейма df создаем двумерный массив
        support_arr = df_support.to_numpy()

        global support_list
        # создаем список
        support_list = support_arr.tolist()

        # создаем столбец сквозной нумерации
        j = 1
        for i in range(0, len(support_list)):
            support_list[i].insert(0, j)
            j += 1

        # создаем список номеров из словаря Котенко для материалов элементов трубопроводов
        list_of_elements = list(range(13, 28))
        list_of_elements_2=list(range(63, 67))
        list_of_elements_3=list(range(71, 73))
        # обьединяем списки
        list_of_elements.extend(list_of_elements_2)
        list_of_elements.extend(list_of_elements_3)
        # создаем датафрейм df_elements в котором только элементы трубопроводов
        df_elements = df_clean[df_clean[columns_names_df[1]].isin(list_of_elements)]

        # удаляем лишние столбцы из  датафрейма df_elements
        df_elements = df_elements.drop(df.columns[[1, 5]], axis=1)
        # cоздаем список из названий столбцов полученного
        columns_names_elements = df_elements.columns.tolist()

        # переименовываем столбец 4 чтобы был как у труб
        df_elements=df_elements.rename(columns={columns_names_elements[3]:columns_names_df[5]})

        #обьединяем датафрейм труб и элементов трубопроводов
        df_elements_new = pd.concat([df_tube_new, df_elements], axis=0)  # объединение по вертикали

        # из датафрейма df создаем двумерный массив
        elements_arr = df_elements_new.to_numpy()

        global elements_list
        # создаем список
        elements_list = elements_arr.tolist()

        # создаем столбец сквозной нумерации
        j = 1
        for i in range(0, len(elements_list)):
            elements_list[i].insert(0, j)
            j += 1


"""""
#Функция создания таблицы здесь не требуется так как мы заполняем существующую в шаблоне
def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table
"""
# Создается оконное приложение
root=Tk()
# Заголовок
root.title('Подготовка актов из спецификации Solidworks')
root["bg"] = "aquamarine4"
# Размер окна
root.geometry('800x500+100+100')
# Иконка в титуле приложения
root.iconbitmap(default=resource_path('res/_brend.ico'))

# Стили
font1 = font.Font(family= "Times New Roman", size=11, weight="bold", slant="roman", underline=False, overstrike=False)
font2 = font.Font(family= "Times New Roman", size=11, weight="normal", slant="roman", underline=False, overstrike=False)

#Текстовая метка
name_form=Label(root, text='Заполните данные шапки акта', font=("Arial", 11, "bold"))
name_form.place(x=20, y=20)

#привязываем переменную name_station к полю ввода названия установки
name_station = StringVar()

#Поле ввода текста
station =Entry(root, font=font1, textvariable=name_station)
station.place(x=20, y=60, width=650)
station.insert(0,'Введите название установки')

#привязываем переменную number_cal к полю ввода номер расчета
number_calc = StringVar()

calc =Entry(root, font=font1, textvariable=number_calc)
calc.place(x=20, y=100, width=650)
calc.insert(0,'Введите номер расчета')

#привязываем переменную company_name к полю ввода название компании
company_name = StringVar()

company =Entry(root, font=font2, textvariable=company_name)
company.place(x=20, y=140, width=650)
company.insert(0,'Введите название компании')

#привязываем переменную obj_name к полю ввода название обьекта
obj_name = StringVar()

obj =Entry(root, font=font2, textvariable=obj_name)
obj.place(x=20, y=180, width=650)
obj.insert(0,'Введите название обьекта')

#привязываем переменную addr_obj к полю ввода адресс обьекта
addr_obj = StringVar()

address =Entry(root, font=font2, textvariable=addr_obj)
address.place(x=20, y=220, width=650)
address.insert(0,'Введите название адреса')

"""""
number =Entry(root, font=font1)
number.place(x=20, y=260, width=650)
number.insert(0,'Введите номер акта')

name =Entry(root, font=font1)
name.place(x=20, y=300, width=650)
name.insert(0,'Введите название акта')

"""

#привязываем переменную dt к полю ввода дата
dt = StringVar()

data =Entry(root, font=font2, textvariable=dt )
data.place(x=20, y=260, width=650)
data.insert(0,'Введите дату')

acttype=Label(root, text='Выберите тип акта', font=("Arial", 11, "bold"))
acttype.place(x=20, y=300)

type_acts=['входного контроля элементов трубопровода', 'входного контроля материалов',
           'проверки чистоты труб', 'монтажа опор под трубопроводы', 'монтажа трубопроводов']
# по умолчанию будет выбран первый элемент из languages
type_var = StringVar(value=type_acts[0])

# Ниспадающий список
type_choies=Combobox(textvariable=type_var, values=type_acts, state="readonly")
type_choies.place(x=20, y=330, width=350)

#Кнопка открытия спецификации
file_button=Button(text='Открыть спец', command=open_table, font=("Arial", 12, "bold"))
file_button.place(x=400, y=20)

#Кнопка открытия файла xlsx c полями шапки
head_button=Button(text='Поля шапки', command=open_head, font=("Arial", 12, "bold"))
head_button.place(x=560, y=20)

#Кнопка создания актов
btn=Button(text='Создать Акт', command=safe_act, font=("Arial", 12, "bold"))
btn.place(x=560, y=320)


# Загрузка шаблона
document = DocxTemplate(resource_path('res\Шаблон_solid.docx'))

tube_list=[]
pad_list=[]
support_list=[]
elements_list=[]

#df_fo_dic = pd.read_excel(resource_path('res\Словарь.xlsx'))
#model_dict=dict(zip(df_fo_dic['id'], df_fo_dic['value']))




root.mainloop()