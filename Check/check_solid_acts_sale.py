from docxtpl import DocxTemplate
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def fill_table_sale(choice:str, col_number:int, elements:list, pad:list, tube:list, support:list, sales:list, table_for_fill:DocxTemplate) -> DocxTemplate:
    # Если выбран тип акта основное оборудование
    if choice == 'входного контроля элементов трубопровода':

        new_elements_list = elements.copy()
        # Заполняется таблица шаблона списком основного оборудования
        for row in new_elements_list:
            row_cells = table_for_fill.add_row().cells
            for i in range(col_number):
                row_cells[i].text = str(row[i])
                # центруем ячейки в таблице
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return table_for_fill

    # Далее аналогично двум другим выборкам
    elif choice == 'входного контроля материалов':

        new_pad_list = pad.copy()
        # Заполняется таблица шаблона списком основного оборудования
        for row in new_pad_list:
            row_cells = table_for_fill.add_row().cells
            for i in range(col_number):
                row_cells[i].text = str(row[i])
                # центруем ячейки в таблице
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return table_for_fill

    elif choice == 'проверки чистоты труб':
        new1_tube_list = tube.copy()
        # Заполняется таблица шаблона списком основного оборудования
        for row in new1_tube_list:
            row_cells = table_for_fill.add_row().cells
            for i in range(col_number):
                row_cells[i].text = str(row[i])
                # центруем ячейки в таблице
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return table_for_fill

    elif choice == 'монтажа опор под трубопроводы':

        new_support_list = support.copy()
        # Заполняется таблица шаблона списком основного оборудования
        for row in new_support_list:
            row_cells = table_for_fill.add_row().cells
            for i in range(col_number):
                row_cells[i].text = str(row[i])
                # центруем ячейки в таблице
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return table_for_fill

    elif choice == 'монтажа трубопроводов':

        new_tube_list = tube.copy()
        # Заполняется таблица шаблона списком основного оборудования
        for row in new_tube_list:
            row_cells = table_for_fill.add_row().cells
            for i in range(col_number):
                row_cells[i].text = str(row[i])
                # центруем ячейки в таблице
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return table_for_fill

    ######
    elif choice == '*закупка*':

        new_sales_list = sales.copy()
        # Заполняется таблица шаблона списком основного оборудования
        for row in new_sales_list:
            row_cells = table_for_fill.add_row().cells
            for i in range(col_number):
                row_cells[i].text = str(row[i])
                # центруем ячейки в таблице
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return table_for_fill

    #####

    # Если выборки не сделано, то таблица в шаблоне наполняется полным списком из спецификации
    else:
        return table_for_fill
